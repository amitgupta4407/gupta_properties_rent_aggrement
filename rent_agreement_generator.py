import streamlit as st
import docx
import re
import io
from datetime import datetime, timedelta
import os
import inflect

# Initialize inflect engine for converting numbers to words
p = inflect.engine()

def format_date_with_suffix(date):
    """Format date with suffix (1st, 2nd, 3rd, etc.)"""
    day = date.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return f"{day}{suffix} {date.strftime('%B %Y')}"

def number_to_words_rupees(number):
    """Convert number to words with 'Rupees' prefix and 'only' suffix"""
    if number == 0:
        return "Rupees Zero only"
    
    words = p.number_to_words(number)
    words = words.replace(',', '')
    words = words.replace('-', ' ')
    words = ' '.join(word.capitalize() for word in words.split())
    return f"Rupees {words} only"

def generate_agreement(template_path, field_values):
    """Generate agreement by replacing placeholders with values"""
    try:
        doc = docx.Document(template_path)
    except Exception as e:
        # If template doesn't exist, create a new document based on the draft.txt
        st.warning(f"Template file not found. Creating a new document based on the draft format.")
        doc = create_document_from_draft(field_values)
        return doc
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        for field_name, value in field_values.items():
            placeholder = f"[[{field_name}]]"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for field_name, value in field_values.items():
                        placeholder = f"[[{field_name}]]"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    return doc

def create_document_from_draft(field_values):
    """Create a document based on the draft.txt format"""
    doc = docx.Document()
    
    # Add title - centered and underlined
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("LEASE DEED")
    title_run.bold = True
    title_run.underline = True
    title_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add execution date - centered
    exec_date_para = doc.add_paragraph()
    exec_date_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    # Add superscript for date suffix (th, st, nd, rd)
    date_text = field_values.get('execution_date', '')
    if date_text:
        day_num = int(''.join(filter(str.isdigit, date_text.split()[0])))
        day_suffix = 'th'
        if day_num % 10 == 1 and day_num != 11:
            day_suffix = 'st'
        elif day_num % 10 == 2 and day_num != 12:
            day_suffix = 'nd'
        elif day_num % 10 == 3 and day_num != 13:
            day_suffix = 'rd'
        
        exec_date_para.add_run(f"This Lease Deed is executed on this {day_num}")
        suffix_run = exec_date_para.add_run(day_suffix)
        suffix_run.font.superscript = True
        exec_date_para.add_run(f" day of {field_values.get('execution_month', '')} '{field_values.get('execution_year', '')}")
    
    # Add BETWEEN section - centered
    between_para = doc.add_paragraph()
    between_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    between_para.add_run("BETWEEN").bold = True
    
    # Add lessor details - with underline for name
    lessor_para = doc.add_paragraph()
    lessor_name_run = lessor_para.add_run(f"MR {field_values.get('lessor_name', '')} S/O {field_values.get('lessor_father_name', '')} {field_values.get('lessor_address', '')}")
    lessor_name_run.underline = True
    lessor_name_run.bold = True
    lessor_para.add_run(" (hereinafter called the Lessor(s)/ Owner(s) which expression unless repugnant to the subject or context thereof shall include his heirs, successors, executors, administrators, legal representatives etc.")
    
    # Add AND section - centered
    and_para = doc.add_paragraph()
    and_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    and_para.add_run("AND").bold = True
    
    # Add lessee details - with underline for name
    lessee_para = doc.add_paragraph()
    lessee_name_run = lessee_para.add_run(f"MR.{field_values.get('lessee_name', '')} S/o {field_values.get('lessee_father_name', '')} R/O {field_values.get('lessee_address', '')}")
    lessee_name_run.underline = True
    lessee_name_run.bold = True
    lessee_para.add_run(" (Hereinafter called the Lessee/ Tenant, which expression unless repugnant to the subject or context thereof shall include its successors, executors, administrators, legal representatives etc.)")
    
    # Add WHEREAS section
    whereas_para = doc.add_paragraph()
    whereas_para.add_run("WHEREAS, ").bold = True
    whereas_para.add_run("the Lessor(s) is the sole and absolute owner and is in actual, physical peaceful possession of the premises at ")
    apartment_run = whereas_para.add_run(f"APARTMENT/ UNIT No. {field_values.get('apartment_unit_no', '')}")
    apartment_run.bold = True
    whereas_para.add_run(" in ")
    tower_run = whereas_para.add_run(f"TOWER NO- {field_values.get('tower_no', '')}")
    tower_run.bold = True
    whereas_para.add_run(", located at ")
    property_run = whereas_para.add_run(f"{field_values.get('property_name', '')}")
    property_run.bold = True
    whereas_para.add_run(", Sector-")
    sector_run = whereas_para.add_run(f"{field_values.get('property_sector', '')}")
    sector_run.bold = True
    whereas_para.add_run(f", {field_values.get('property_location', '')}")
    whereas_para.add_run(" hereinafter referred to as the \"Said Premises\").")
    
    # Add AND Whereas section
    and_whereas_para = doc.add_paragraph()
    and_whereas_para.add_run("AND Whereas ").bold = True
    and_whereas_para.add_run("on request of the lessee, the Lessor aforesaid has agreed to let-out the said premises and 'lessee' after his satisfaction has agreed to take on lease the 'said premises' for RESIDENTIAL purpose and shall not use it for any other purpose and whereas the lessee has agreed to execute and sign this deed of rent agreement as per terms and conditions mentioned below :-")
    
    # Add NOW THIS LEASE DEED WITNESSETH AS UNDER section
    now_para = doc.add_paragraph()
    now_para.add_run("NOW THIS LEASE DEED WITNESSETH AS UNDER:").bold = True
    
    # Format clauses with proper spacing and bold for important values
    clauses = [
        # Rent amount
        {"text": "That the rent for the demised property and fittings provided therein payable by the tenant of the owner shall be Rs. [[rent_amount]]/- ( [[rent_amount_words]] excluding maintenance) which shall be directly payable by tenant & will be applicable from 01-May-[[execution_year]].", 
         "bold_fields": ["rent_amount", "execution_year"]},
        
        # Security deposit
        {"text": "The tenant shall deposit Rs. [[security_deposit]]/- ([[security_deposit_words]]) through cheque/online as a security deposit to the owner, free of interest. This deposit shall be refunded to the tenant upon the expiry of the lease, after deducting any arrears on rent, electricity, water dues, any damage, and cleaning costs of the flat, if any. This clause will apply after the flat is vacant and will not be adjustable against the rental.",
         "bold_fields": ["security_deposit"]},
        
        # Lease period
        {"text": "That the lease is for a period of [[lease_period_months]] months w.e.f. 1st May [[execution_year]] to 31st March [[next_year]].",
         "bold_fields": ["lease_period_months", "execution_year", "next_year"]},
        
        # Payment due day
        {"text": "The monthly rent shall be payable on or before [[payment_due_day]]th of each English Calendar month by Online transfer directly to owners account.",
         "bold_fields": ["on or before [[payment_due_day]]th"]},
        
        # Rent increase
        {"text": "That after the expiry of [[lease_period_months]] months period the rent shall be increased by [[rent_increase_percentage]]% if the tenancy shall be continued and tenant to deposit rent by [[rent_increase_percentage]]% increase after [[lease_period_months]]th month.",
         "bold_fields": ["lease_period_months", "increased by [[rent_increase_percentage]]%"]},
        
        # Electricity charges
        {"text": "That the electricity charges, Gas(IGL) and water charges and other society charges shall be paid by the lessee directly to the Estate Office or concerned authority as per the meter installed therein. That the monthly Society maintenance charges shall be paid by the lessee directly to the Estate office or concerned authority before the due date as per monthly invoice.",
         "bold_fields": []},
        
        # Subletting
        {"text": "That the Lessee or his occupants shall not sub-let, assign or part with possession of the said premises in question or any portion thereof in any manner whatsoever.",
         "bold_fields": []},
        
        # Inspection
        {"text": "That the Lessee shall permit the Lessor or its agents/authorized person to enter the deemed premises to inspect, click photos of flat premises and view the state and condition thereof at reasonable times during the tenancy period, but with an advance notice.",
         "bold_fields": []},
        
        # Structural changes
        {"text": "That the lessee shall not carry out any permanent or temporary structural additions or alterations to the building layout.",
         "bold_fields": []},
        
        # Repairs
        {"text": "That the lessor shall effect all major repairs such as major leakage in water pipes or major structural cracks etc. at his / her own cost immediately upon such defects being notified to him / her by the lessee and all minor repairs will be borne by the lessee",
         "bold_fields": []},
        
        # Security deposit refund
        {"text": "The interest free deposit shall be refundable on termination of lease subject to handing over of vacant physical possession and all fixtures & fittings in working conditions, by the lessor after adjustment of Water, Electricity Charges, Gas, cleaning Etc. if outstanding in any case.",
         "bold_fields": []},
        
        # Notice period
        {"text": "That the Lessor / Lessee shall be at liberty to revoke the present lease at any time by serving [[notice_period_months]] month notice in writing or by paying One month rent in lieu thereof.",
         "bold_fields": ["notice_period_months"]},
        
        # Combustible goods
        {"text": "That the lessee shall not store in the demised premises or any part thereof any such goods of combustible or explosive nature, provided that nothing contained in this sub-clause shall apply to the storage of kerosene, lanterns etc. kept for day-to-day use.",
         "bold_fields": []},
        
        # Expiry terms
        {"text": "On the expiry of the terms of this license, the Licensee shall remove itself, its servants and goods from the said FLAT without demur and without raising any objection of any sort or kind whatsoever and shall not claim any tenancy rights in the said FLAT.",
         "bold_fields": []},
        
        # Rules compliance
        {"text": "That the Lessee shall comply with all the rules & regulations of the local authorities and Society, whatsoever with relation to the use and occupation of the said premises.",
         "bold_fields": []},
        
        # Residential use
        {"text": "That the Tenant shall use the said premises only for residential purposes of self and dependent family and not for any other purposes.",
         "bold_fields": []},
        
        # No subletting
        {"text": "That the tenant/second party shall not sub-let the said premises of any portion thereof to any persons or persons under any circumstances",
         "bold_fields": []},
        
        # Antisocial activities
        {"text": "The Lessee shall not carry out any acts or activities which are obnoxious, antisocial, illegal or prejudicial to the norms of decency or etiquette or society by laws which cause a nuisance to the other members of the society in the building.",
         "bold_fields": []},
        
        # Early termination
        {"text": "If Lessee terminate the lease before 6 month of start date, the security amount will not be refundable.",
         "bold_fields": []},
        
        # Possession return
        {"text": "The Lessee agrees to deliver vacant and peaceful possession of the said FLAT on the expiry of this agreement in good and clean condition as it was when the lessee obtained possession unless extended for a further period of 11 month by mutual consent.",
         "bold_fields": []},
        
        # Property sale
        {"text": "In the event that the Owner decides to sell the property, the Tenant agrees to accommodate reasonable requests for property viewings and inspections by potential buyers, provided that the Tenant is given at least 24 hours' notice",
         "bold_fields": []},
        
        # Police verification
        {"text": "That in accordance with the requirement of law, the lessee shall get police verification done and produce document/s for satisfaction of the said authority.",
         "bold_fields": []},
        
        # Property type
        {"text": "That the Lessor is providing furnished flat consisting of [[property_type]].",
         "bold_fields": ["property_type"]},
        
        # Car parks
        {"text": "Lessor is providing [[car_parks]] car park in the society premises for exclusive use of the tenant [[lessee_name]] & his family and not to sublease the parking further.",
         "bold_fields": ["car_parks", "lessee_name"]}
    ]
    
    # Process each clause
    for i, clause_data in enumerate(clauses, 1):
        clause_para = doc.add_paragraph()
        clause_para.paragraph_format.left_indent = docx.shared.Pt(12)
        clause_para.paragraph_format.first_line_indent = docx.shared.Pt(-12)
        
        # Add number with less spacing
        number_run = clause_para.add_run(f"{i}.")
        number_run.bold = True
        
        # Add space after number
        clause_para.add_run(" ")
        
        # Get the clause text and replace placeholders with values
        text = clause_data["text"]
        bold_fields = clause_data["bold_fields"]
        
        # Replace placeholders with actual values
        for field_name, value in field_values.items():
            placeholder = f"[[{field_name}]]"
            if placeholder in text:
                # Convert value to string to avoid type errors
                text = text.replace(placeholder, str(value))
        
        # Special case for next year in clause 3
        if i == 3:
            next_year = int(str(field_values.get('execution_year', 2025))) + 1
            text = text.replace("[[next_year]]", str(next_year))
        
        # Add the text with bold parts
        if bold_fields:
            # Split text into parts to bold specific sections
            current_text = text
            for bold_field in bold_fields:
                # Handle special cases
                if bold_field == "on or before [[payment_due_day]]th":
                    bold_text = f"on or before {field_values.get('payment_due_day', '5')}th"
                    parts = current_text.split(bold_text)
                    if len(parts) > 1:
                        clause_para.add_run(parts[0])
                        bold_run = clause_para.add_run(bold_text)
                        bold_run.bold = True
                        current_text = parts[1]
                elif bold_field == "increased by [[rent_increase_percentage]]%":
                    bold_text = f"increased by {field_values.get('rent_increase_percentage', '10')}%"
                    parts = current_text.split(bold_text)
                    if len(parts) > 1:
                        clause_para.add_run(parts[0])
                        bold_run = clause_para.add_run(bold_text)
                        bold_run.bold = True
                        current_text = parts[1]
                else:
                    # Regular field replacement
                    field_value = str(field_values.get(bold_field, ""))
                    if field_value:
                        parts = current_text.split(field_value)
                        if len(parts) > 1:
                            clause_para.add_run(parts[0])
                            bold_run = clause_para.add_run(field_value)
                            bold_run.bold = True
                            current_text = "".join(parts[1:])
            
            # Add any remaining text
            clause_para.add_run(current_text)
        else:
            # No bold parts, just add the text
            clause_para.add_run(text)
    
    # Add signature section
    doc.add_paragraph("\nIN WITNESS WHEREOF, the parties have placed their respective hands and signed this Lease Deed on this date     Day of________, in the presence of the following witnesses.\n\n\n")
    
    # Add signature lines
    sig_para = doc.add_paragraph()
    sig_para.add_run("        (LESSOR/FIRST PARTY)").bold = True
    sig_para.add_run("                                                         ").bold = True
    sig_para.add_run("(LESSEE / SECOND PATY)").bold = True
    
    # Add witness section
    witness_para = doc.add_paragraph("\n\n\nWitness 1.  _________________________\t\t Witness 2.  __________________________")
    doc.add_paragraph("\t      _________________________\t\t                     __________________________")
    doc.add_paragraph("\t      _________________________\t\t\t       __________________________")
    
    # Add Annexure section
    doc.add_paragraph("\n\n")
    section = doc.add_section(docx.enum.section.WD_SECTION_START.NEW_PAGE)
    annexure_para = doc.add_paragraph()
    annexure_run = annexure_para.add_run("Annexure")
    annexure_run.bold = True
    # Remove the blue color to keep it black
    
    # Add LIST OF FURNITURE & FIXTURES heading - centered
    fixtures_para = doc.add_paragraph()
    fixtures_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    fixtures_run = fixtures_para.add_run("LIST OF FURNITURE & FIXTURES")
    fixtures_run.bold = True
    fixtures_run.underline = True
    
    # Create table for furniture and fixtures
    table = doc.add_table(rows=14, cols=3)
    table.style = 'Table Grid'
    
    # Set column widths
    for cell in table.columns[0].cells:
        cell.width = docx.shared.Inches(0.5)
    for cell in table.columns[1].cells:
        cell.width = docx.shared.Inches(2.0)
    for cell in table.columns[2].cells:
        cell.width = docx.shared.Inches(3.0)
    
    # Add headers
    headers = ["S.NO.", "ITEM", "DESCRIPTION"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add furniture items
    furniture_items = [
        ["1", "Prepaid Metering System", ""],
        ["2", "Ceiling Fans", "6 Nos."],
        ["3", "Tube lights/ Wall Lights/Ceiling Lights", "6 LED Tubelights"],
        ["4", "Fan Regulators", "In all rooms"],
        ["5", "Electric Bell", "1 Nos."],
        ["6", "Geyser", "1 in master bed Toilet"],
        ["7", "Electric Auto Clean Chimney", "1 in the Kitchen"],
        ["8", "Mirror", "2 in the both Toilet"],
        ["9", "Modular Wood Work cabinet", "In the Kitchen"],
        ["10", "Fixed Almirah", "Both Bedrooms"],
        ["11", "Keys", "Single key of every door"],
        ["12", "Other Facilities", "Club Facilities provided by builder. Usage on chargeable basis"]
    ]
    
    # Add furniture items to table
    for i, item in enumerate(furniture_items):
        row = table.rows[i+1]  # Skip header row
        for j, text in enumerate(item):
            cell = row.cells[j]
            cell.text = text
    
    # Add Lessor and Lessee signature lines at the bottom
    doc.add_paragraph("\n\n\n")
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    
    # Set column widths
    sig_table.columns[0].width = docx.shared.Inches(3.0)
    sig_table.columns[1].width = docx.shared.Inches(3.0)
    
    # Add Lessor and Lessee text
    lessor_cell = sig_table.cell(0, 0)
    lessor_cell.text = "Lessor"
    
    lessee_cell = sig_table.cell(0, 1)
    lessee_cell.text = "Lessee"
    lessee_cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    
    return doc

def main():
    st.title("Gupta Properties: Rent Agreement Generator")
    st.subheader("Based on Lease Deed Template")
    
    # Template path
    template_path = "agreement_template.docx"
    
    # Create tabs for better organization
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "Basic Information", 
        "Financial Details", 
        "Lease Period", 
        "Property Details",
        "Furniture & Fixtures"
    ])
    
    field_values = {}
    validation_errors = []
    
    with tab1:
        st.subheader("Parties & Property Information")
        
        # Execution date
        execution_date = st.date_input("Execution Date", key="execution_date")
        field_values["execution_date"] = format_date_with_suffix(execution_date)
        field_values["execution_month"] = execution_date.strftime("%B")
        field_values["execution_year"] = execution_date.strftime("%Y")
        
        # Lessor details
        st.markdown("### Lessor (Owner) Details")
        lessor_name = st.text_input("Lessor Name", key="lessor_name")
        if not lessor_name:
            validation_errors.append("Lessor name is required")
        field_values["lessor_name"] = lessor_name
        
        lessor_father_name = st.text_input("Lessor's Father Name (S/O)", key="lessor_father")
        if not lessor_father_name:
            validation_errors.append("Lessor's father name is required")
        field_values["lessor_father_name"] = lessor_father_name
        
        lessor_address = st.text_area("Lessor Address", key="lessor_address")
        if not lessor_address:
            validation_errors.append("Lessor address is required")
        field_values["lessor_address"] = lessor_address
        
        # Lessee details
        st.markdown("### Lessee (Tenant) Details")
        lessee_name = st.text_input("Lessee Name", key="lessee_name")
        if not lessee_name:
            validation_errors.append("Lessee name is required")
        field_values["lessee_name"] = lessee_name
        
        lessee_father_name = st.text_input("Lessee's Father Name (S/O)", key="lessee_father")
        if not lessee_father_name:
            validation_errors.append("Lessee's father name is required")
        field_values["lessee_father_name"] = lessee_father_name
        
        lessee_address = st.text_area("Lessee Address (with PIN code)", key="lessee_address")
        if not lessee_address:
            validation_errors.append("Lessee address is required")
        field_values["lessee_address"] = lessee_address
    
    with tab2:
        st.subheader("Financial Information")
        
        # Rent amount
        rent_amount = st.number_input("Monthly Rent (₹)", min_value=0, value=43000, key="rent")
        field_values["rent_amount_numeric"] = f"Rs. {rent_amount:,}/-"
        field_values["rent_amount_words"] = number_to_words_rupees(rent_amount).replace("Rupees ", "")
        
        # Security deposit
        security_deposit = st.number_input("Security Deposit (₹)", min_value=0, value=rent_amount*2, key="deposit")
        field_values["security_deposit_numeric"] = f"Rs. {security_deposit:,}/-"
        field_values["security_deposit_words"] = number_to_words_rupees(security_deposit).replace("Rupees ", "")
        
        # Payment details
        col1, col2 = st.columns(2)
        with col1:
            payment_due_day = st.number_input(
                "Payment Due Day", min_value=1, max_value=31, value=5, key="due_day",
                help="Day of month when rent is due (e.g., 5 means rent is due by 5th of each month)"
            )
            field_values["payment_due_day"] = payment_due_day
        with col2:
            rent_increase = st.number_input(
                "Rent Increase Percentage", min_value=0, max_value=100, value=10, key="increase",
                help="Percentage by which rent will increase after the lease period"
            )
            field_values["rent_increase_percentage"] = rent_increase
    
    with tab3:
        st.subheader("Lease Period")
        
        # Lease start and end dates
        col1, col2 = st.columns(2)
        with col1:
            lease_start = st.date_input(
                "Lease Start Date", key="lease_start",
                help="First day of the lease period"
            )
            field_values["lease_start_date"] = format_date_with_suffix(lease_start)
        
        with col2:
            # Calculate default end date (11 months from start date)
            default_end_date = lease_start + timedelta(days=11*30)  # Approximate 11 months
            lease_end = st.date_input(
                "Lease End Date", 
                value=default_end_date,
                key="lease_end",
                help="Last day of the lease period (typically 11 months from start date)"
            )
            field_values["lease_end_date"] = format_date_with_suffix(lease_end)
        
        # Calculate lease period in months
        lease_period_months = 11  # Default
        if lease_start and lease_end:
            # Calculate difference in months (approximate)
            delta = lease_end - lease_start
            lease_period_months = round(delta.days / 30)
        
        # Lease period
        lease_period = st.number_input(
            "Lease Period (months)", 
            min_value=1, 
            max_value=60,
            value=lease_period_months,
            help="Duration of lease in months (typically 11 months in India)",
            key="lease_period"
        )
        field_values["lease_period_months"] = lease_period
        
        # Notice period
        notice_period = st.selectbox(
            "Notice Period", 
            options=["one", "two", "three"],
            index=1,  # Default to "two"
            help="Notice period required before terminating lease",
            key="notice_period"
        )
        field_values["notice_period_months"] = notice_period
    
    with tab4:
        st.subheader("Property Details")
        
        # Property details
        col1, col2 = st.columns(2)
        with col1:
            apartment_unit = st.text_input(
                "Apartment/Unit Number", 
                key="apartment_unit",
                help="Apartment or unit number"
            )
            if not apartment_unit:
                validation_errors.append("Apartment/Unit number is required")
            field_values["apartment_unit_no"] = apartment_unit
            
            tower_no = st.text_input(
                "Tower Number", 
                key="tower_no",
                help="Tower or building number"
            )
            if not tower_no:
                validation_errors.append("Tower number is required")
            field_values["tower_no"] = tower_no
        
        with col2:
            property_name = st.text_input(
                "Property Name", 
                key="property_name",
                help="Name of the housing society or complex"
            )
            if not property_name:
                validation_errors.append("Property name is required")
            field_values["property_name"] = property_name
            
            property_sector = st.text_input(
                "Sector", 
                key="property_sector",
                help="Sector or area number"
            )
            if not property_sector:
                validation_errors.append("Property sector is required")
            field_values["property_sector"] = property_sector
        
        property_location = st.text_input(
            "Location", 
            key="property_location",
            help="City and state (e.g., Noida, UP)"
        )
        if not property_location:
            validation_errors.append("Property location is required")
        field_values["property_location"] = property_location
        
        property_type = st.text_input(
            "Property Type", 
            value="3 BHK + Study",
            key="property_type",
            help="Type of property (e.g., 2 BHK, 3 BHK + Study)"
        )
        field_values["property_type"] = property_type
        
        car_parks = st.number_input(
            "Number of Car Parks", 
            min_value=0, 
            value=2,
            key="car_parks",
            help="Number of car parking spaces included"
        )
        field_values["car_parks"] = car_parks
    
    with tab5:
        st.subheader("Furniture & Fixtures")
        
        # Furniture and fixtures
        ceiling_fans = st.number_input(
            "Ceiling Fans", 
            min_value=0, 
            value=6,
            key="ceiling_fans"
        )
        field_values["ceiling_fans"] = ceiling_fans
        
        tube_lights = st.text_input(
            "Tube Lights/Wall Lights/Ceiling Lights", 
            value="6 LED Tubelights",
            key="tube_lights"
        )
        field_values["tube_lights"] = tube_lights
        
        fan_regulators = st.text_input(
            "Fan Regulators", 
            value="In all rooms",
            key="fan_regulators"
        )
        field_values["fan_regulators"] = fan_regulators
        
        electric_bell = st.number_input(
            "Electric Bell", 
            min_value=0, 
            value=1,
            key="electric_bell"
        )
        field_values["electric_bell"] = electric_bell
        
        geyser = st.text_input(
            "Geyser", 
            value="1 in master bed Toilet",
            key="geyser"
        )
        field_values["geyser"] = geyser
        
        chimney = st.text_input(
            "Electric Auto Clean Chimney", 
            value="1 in the Kitchen",
            key="chimney"
        )
        field_values["chimney"] = chimney
        
        mirrors = st.text_input(
            "Mirrors", 
            value="2 in the both Toilet",
            key="mirrors"
        )
        field_values["mirrors"] = mirrors
        
        modular_woodwork = st.text_input(
            "Modular Wood Work Cabinet", 
            value="In the Kitchen",
            key="modular_woodwork"
        )
        field_values["modular_woodwork"] = modular_woodwork
        
        fixed_almirah = st.text_input(
            "Fixed Almirah", 
            value="Both Bedrooms",
            key="fixed_almirah"
        )
        field_values["fixed_almirah"] = fixed_almirah
        
        keys = st.text_input(
            "Keys", 
            value="Single key of every door",
            key="keys"
        )
        field_values["keys"] = keys
    
    # Generate button
    if st.button("Generate Agreement"):
        # Validate all required fields
        if validation_errors:
            for error in validation_errors:
                st.error(error)
        else:
            # Generate document
            try:
                doc = generate_agreement(template_path, field_values)
                
                # Save to BytesIO object
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                # Provide download button
                st.download_button(
                    label="Download Agreement",
                    data=buffer,
                    file_name="rent_agreement.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("Agreement generated successfully!")
            except Exception as e:
                st.error(f"Error generating agreement: {str(e)}")

if __name__ == "__main__":
    main()









